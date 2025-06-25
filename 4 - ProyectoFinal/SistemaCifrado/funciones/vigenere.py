"""
vigenere.py - Módulo de cifrado y descifrado de archivos usando el algoritmo Vigenère.
Soporta archivos de texto (.txt), documentos Word (.docx) y PDF (.pdf).
Incluye funciones para leer archivos, guardar resultados, y aplicar el cifrado y descifrado.
"""

import os
import docx
from PyPDF2 import PdfReader
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter

def leerArchivo(ruta):
    """
    Lee el contenido de un archivo (.txt, .pdf, .docx) y lo convierte a texto plano.

    Parámetros:
        ruta (str): Ruta del archivo a leer.

    Retorna:
        str: Contenido del archivo como texto plano.

    Lanza:
        ValueError: Si el formato del archivo no es compatible.
    """
    ext = os.path.splitext(ruta)[1].lower()

    if ext == '.txt':
        with open(ruta, 'r', encoding='utf-8') as f:
            return f.read()

    elif ext == '.pdf':
        reader = PdfReader(ruta)
        texto = ""
        for pagina in reader.pages:
            texto += pagina.extract_text() or ""
        return texto

    elif ext == '.docx':
        doc = docx.Document(ruta)
        return "\n".join(p.text for p in doc.paragraphs)

    else:
        raise ValueError("❌ Formato de archivo no soportado para lectura.")

def guardarArchivo(texto, ruta):
    """
    Guarda texto plano en un archivo con formato .txt, .pdf o .docx.

    Parámetros:
        texto (str): Texto a escribir.
        ruta (str): Ruta donde se guardará el archivo.

    Lanza:
        ValueError: Si el formato del archivo no es compatible.
    """
    ext = os.path.splitext(ruta)[1].lower()

    if ext == '.txt':
        with open(ruta, 'w', encoding='utf-8') as f:
            f.write(texto)

    elif ext == '.pdf':
        c = canvas.Canvas(ruta, pagesize=letter)
        width, height = letter
        y = height - 40
        for linea in texto.split('\n'):
            c.drawString(40, y, linea[:100])
            y -= 15
            if y < 40:
                c.showPage()
                y = height - 40
        c.save()

    elif ext == '.docx':
        doc = docx.Document()
        for linea in texto.split('\n'):
            doc.add_paragraph(linea)
        doc.save(ruta)

    else:
        raise ValueError("❌ Formato de archivo no soportado para escritura.")

def cifrarVigenere(texto, clave):
    """
    Cifra un texto plano usando el algoritmo Vigenère.

    Parámetros:
        texto (str): Texto original a cifrar.
        clave (str): Clave secreta para el cifrado (solo letras).

    Retorna:
        str: Texto cifrado.
    """
    resultado = ""
    clave = clave.upper()
    indiceClave = 0

    for letra in texto:
        if letra.isalpha():
            offset = 65 if letra.isupper() else 97
            k = ord(clave[indiceClave % len(clave)]) - 65
            nuevaLetra = chr((ord(letra) - offset + k) % 26 + offset)
            resultado += nuevaLetra
            indiceClave += 1
        else:
            resultado += letra

    return resultado

def descifrarVigenere(texto, clave):
    """
    Descifra un texto previamente cifrado con Vigenère.

    Parámetros:
        texto (str): Texto cifrado.
        clave (str): Clave utilizada originalmente para cifrar.

    Retorna:
        str: Texto descifrado.
    """
    resultado = ""
    clave = clave.upper()
    indiceClave = 0

    for letra in texto:
        if letra.isalpha():
            offset = 65 if letra.isupper() else 97
            k = ord(clave[indiceClave % len(clave)]) - 65
            nuevaLetra = chr((ord(letra) - offset - k) % 26 + offset)
            resultado += nuevaLetra
            indiceClave += 1
        else:
            resultado += letra

    return resultado
