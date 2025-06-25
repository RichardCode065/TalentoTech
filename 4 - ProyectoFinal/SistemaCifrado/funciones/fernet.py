"""
fernet.py - Módulo de cifrado y descifrado de archivos usando Fernet.
Soporta archivos .txt, .pdf y .docx.
Incluye funciones para generación/carga de claves, lectura/escritura de archivos,
y operaciones de cifrado y descifrado.
"""

import os
import docx
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from docx import Document
from cryptography.fernet import Fernet
from PyPDF2 import PdfReader

def generarClave(rutaDestino):
    """Genera una clave Fernet y la guarda en un archivo."""
    clave = Fernet.generate_key()
    with open(rutaDestino, 'wb') as archivoClave:
        archivoClave.write(clave)
    return clave

def cargarClave(ruta):
    """Carga una clave Fernet desde un archivo."""
    with open(ruta, 'rb') as archivo:
        return archivo.read()

def leerContenido(ruta):
    """Lee el contenido de un archivo (.txt, .pdf, .docx) y lo retorna como texto plano."""
    extension = os.path.splitext(ruta)[1].lower()
    if extension == '.txt':
        with open(ruta, 'r', encoding='utf-8') as f:
            return f.read()
    elif extension == '.pdf':
        reader = PdfReader(ruta)
        texto = ""
        for pagina in reader.pages:
            texto += pagina.extract_text() or ""
        return texto
    elif extension == '.docx':
        doc = Document(ruta)
        return "\n".join(p.text for p in doc.paragraphs)
    else:
        raise ValueError("Tipo de archivo no soportado para lectura")

def escribirContenido(ruta, contenido):
    """Escribe contenido plano en un archivo con extensión .txt, .pdf o .docx según su extensión."""
    extension = os.path.splitext(ruta)[1].lower()

    if extension == '.txt':
        with open(ruta, 'w', encoding='utf-8') as f:
            f.write(contenido)

    elif extension == '.pdf':
        c = canvas.Canvas(ruta, pagesize=letter)
        width, height = letter
        y = height - 40
        for linea in contenido.split('\n'):
            c.drawString(40, y, linea[:100])
            y -= 15
            if y < 40:
                c.showPage()
                y = height - 40
        c.save()

    elif extension == '.docx':
        doc = Document()
        for linea in contenido.split('\n'):
            doc.add_paragraph(linea)
        doc.save(ruta)

    else:
        raise ValueError("Tipo de archivo no soportado para escritura")

def cifrarArchivo(rutaEntrada, clave, rutaSalida):
    """
    Cifra el contenido de un archivo y lo guarda como binario,
    incluyendo la extensión original como encabezado.
    """
    fernet = Fernet(clave)
    extension = os.path.splitext(rutaEntrada)[1].lower()
    texto = leerContenido(rutaEntrada)

    # Incluir extensión original como encabezado (ejemplo: "EXT:.pdf\n")
    datosConExtension = f"EXT:{extension}\n{texto}"
    datosCifrados = fernet.encrypt(datosConExtension.encode('utf-8'))

    with open(rutaSalida, 'wb') as f:
        f.write(datosCifrados)

    return rutaSalida

def descifrarArchivo(rutaEntrada, clave, rutaSalidaSinExtension):
    """
    Descifra un archivo cifrado con Fernet y lo reconstruye con su extensión original.
    Retorna la ruta del archivo creado o un mensaje de error.
    """
    if not os.path.exists(rutaEntrada):
        return "❌ El archivo cifrado no existe."

    try:
        fernet = Fernet(clave)
        with open(rutaEntrada, 'rb') as f:
            datosCifrados = f.read()

        datosDescifrados = fernet.decrypt(datosCifrados).decode('utf-8')

        # Extraer la extensión original
        if not datosDescifrados.startswith("EXT:"):
            return "❌ No se encontró la extensión original en el archivo cifrado."

        linea, contenido = datosDescifrados.split("\n", 1)
        extension = linea.replace("EXT:", "").strip()

        rutaSalida = rutaSalidaSinExtension + extension
        escribirContenido(rutaSalida, contenido)
        return rutaSalida

    except Exception as e:
        return f"❌ Error durante el descifrado: {str(e)}"
